# deploy: 2026-07-14  ← VERSÃO CORRIGIDA
# =============================================================
# RURALCAIXA — Módulo de Contratos Rurais
# Arquivo: contratos_api.py
# Stack: FastAPI + psycopg2
#
# CORREÇÕES APLICADAS:
#   1. percentual_outorgante/outorgado agora são Optional[float]
#      — condomínio e comodato não exigem esses campos
#   2. Validator aceita None para tipos sem divisão percentual
#   3. Tolerância de ±0.01 na soma dos percentuais (evita erro de float)
#   4. Condomínio usa área em hectares — percentual calculado pelo sistema
# =============================================================

from fastapi import APIRouter, HTTPException, Request, Response
from pydantic import BaseModel, Field, ConfigDict, validator, model_validator
from typing import Optional
from datetime import datetime, timedelta, date
import psycopg2
import psycopg2.extras
import random
import hashlib
import json
import os
import io

router = APIRouter(prefix="/contratos", tags=["Contratos Rurais"])

# Tipos que NÃO usam divisão percentual outorgante/outorgado
TIPOS_SEM_PERCENTUAL = {"condominio", "comodato"}

def get_db():
    return psycopg2.connect(
        os.getenv("DATABASE_URL")
        or "postgresql://postgres:tkyfcRsbrZuuHoThKgjuTiZWYVXOTdOX@gondola.proxy.rlwy.net:53900/railway",
        cursor_factory=psycopg2.extras.RealDictCursor
    )

def log_auditoria(cur, contrato_id, evento, descricao="", ip=None, metadata=None):
    cur.execute(
        """INSERT INTO auditoria_contratos (contrato_id, evento, descricao, ip, metadata)
           VALUES (%s, %s, %s, %s, %s)""",
        (contrato_id, evento, descricao, ip, json.dumps(metadata or {}))
    )

def gerar_otp():
    return str(random.randint(100000, 999999))

def hash_otp(otp: str) -> str:
    return hashlib.sha256(otp.encode()).hexdigest()


# -------------------------------------------------------------
# MODELS
# -------------------------------------------------------------

class ParceiroExterno(BaseModel):
    nome: str
    tipo_documento: str   # CPF ou CNPJ
    documento: str
    telefone: Optional[str] = None
    email: Optional[str] = None

# Mapeamento de nomes em PT-BR para os valores internos aceitos pelo banco
_TIPO_MAP = {
    "agricola": "agricola",
    "pecuaria": "pecuaria",
    "agroindustrial": "agroindustrial",
    "extrativa": "extrativa",
    "condominio": "condominio",
    "arrendamento": "arrendamento",
    "comodato": "comodato",
    "compra_venda": "compra_venda",
    "integracao_agroindustrial": "integracao_agroindustrial",
    "parceria": "pecuaria",
    "parceria rural": "pecuaria",
    "arrendamento rural": "arrendamento",
    "comodato rural": "comodato",
    "condomínio rural": "condominio",
    "condominio rural": "condominio",
    "integração agroindustrial": "integracao_agroindustrial",
    "integracao agroindustrial": "integracao_agroindustrial",
}

class ContratoCreate(BaseModel):
    model_config = ConfigDict(populate_by_name=True)

    # aceita tanto fazenda_id (nome interno) quanto imovel_id (nome que o
    # frontend Next.js envia)
    fazenda_id: int = Field(..., alias="fazenda_id")
    imovel_id: Optional[int] = Field(None, alias="imovel_id")

    tipo: str
    outorgante_socio_id: Optional[int] = None
    outorgante_externo: Optional[ParceiroExterno] = None
    outorgado_socio_id: Optional[int] = None
    outorgado_externo: Optional[ParceiroExterno] = None
    data_inicio: str                   # YYYY-MM-DD
    data_fim: Optional[str] = None     # em branco = prazo indeterminado (comum em condomínio/comodato)
    # ✅ CORREÇÃO 1: Optional — condomínio e comodato não usam esses campos
    percentual_outorgante: Optional[float] = None
    percentual_outorgado: Optional[float] = None
    frequencia_pagamento: str = "safra"
    area_parceria_hectares: Optional[float] = None
    clausulas_adicionais: Optional[dict] = {}

    @model_validator(mode="before")
    @classmethod
    def resolver_fazenda_id(cls, values):
        """Se o frontend enviar imovel_id mas não fazenda_id, usa imovel_id."""
        if isinstance(values, dict):
            if not values.get("fazenda_id") and values.get("imovel_id"):
                values["fazenda_id"] = values["imovel_id"]
        return values

    @validator("tipo")
    def tipo_valido(cls, v):
        normalizado = v.lower().strip()
        mapeado = _TIPO_MAP.get(normalizado)
        if not mapeado:
            validos = sorted(set(_TIPO_MAP.values()))
            raise ValueError(
                f"tipo inválido: '{v}'. Valores aceitos: {validos}. "
                f"Aliases PT-BR aceitos: parceria, arrendamento rural, "
                f"comodato rural, condomínio rural, integração agroindustrial."
            )
        return mapeado

    # ✅ CORREÇÃO 2: Validator aceita None para condomínio e comodato
    @validator("percentual_outorgado", pre=True, always=True)
    def validar_percentuais(cls, v, values):
        tipo = values.get("tipo", "")

        # Condomínio e Comodato não usam percentual outorgante/outorgado
        if tipo in TIPOS_SEM_PERCENTUAL:
            return v  # aceita None sem validação

        # Para os demais tipos, os percentuais são obrigatórios
        if v is None:
            raise ValueError(
                f"percentual_outorgado é obrigatório para contratos do tipo '{tipo}'."
            )
        ote = values.get("percentual_outorgante")
        if ote is None:
            raise ValueError(
                f"percentual_outorgante é obrigatório para contratos do tipo '{tipo}'."
            )

        # ✅ CORREÇÃO 3: Tolerância de ±0.01 para evitar erro de ponto flutuante
        soma = float(ote) + float(v)
        if abs(soma - 100.0) > 0.01:
            raise ValueError(
                f"percentual_outorgante ({ote}) + percentual_outorgado ({v}) = {soma:.2f}. "
                f"A soma deve ser 100%."
            )
        return v

class AssinarRequest(BaseModel):
    papel: str      # outorgante | outorgado
    otp: str
    geolocalizacao: Optional[dict] = None


# -------------------------------------------------------------
# GET /contratos
# -------------------------------------------------------------
_STATIC_PATHS = {"acerto", "novo", "resumo", "relatorio"}

@router.get("/")
def listar_contratos(
    fazenda_id: Optional[int] = None,
    status: Optional[str] = None,
    tipo: Optional[str] = None
):
    conn = get_db()
    try:
        cur = conn.cursor()
        wheres = []
        params = []

        if fazenda_id:
            wheres.append("fazenda_id = %s")
            params.append(fazenda_id)
        if status:
            wheres.append("status = %s")
            params.append(status)
        if tipo:
            wheres.append("tipo = %s")
            params.append(tipo)

        sql = "SELECT * FROM vw_contratos_resumo"
        if wheres:
            sql += " WHERE " + " AND ".join(wheres)
        sql += " ORDER BY criado_em DESC"

        cur.execute(sql, params)
        rows = cur.fetchall()
        return {"data": [dict(r) for r in rows], "total": len(rows)}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()


# -------------------------------------------------------------
# GET /contratos/{id}
# -------------------------------------------------------------
@router.get("/{contrato_id}")
def detalhe_contrato(contrato_id: str):
    if contrato_id in _STATIC_PATHS:
        raise HTTPException(status_code=404, detail="Rota de frontend — não é um contrato.")
    conn = get_db()
    try:
        cur = conn.cursor()
        cur.execute("SELECT * FROM vw_contratos_resumo WHERE id = %s", (contrato_id,))
        contrato = cur.fetchone()
        if not contrato:
            raise HTTPException(status_code=404, detail="Contrato não encontrado")

        cur.execute("""
            SELECT a.id, a.papel, a.status, a.assinado_em, a.visualizado_em,
                   a.ip_assinatura, a.geolocalizacao,
                   s.nome  AS socio_nome,
                   pe.nome AS parceiro_nome
            FROM assinaturas a
            LEFT JOIN produtores s              ON s.id  = a.socio_id
            LEFT JOIN parceiros_externos pe ON pe.id = a.parceiro_externo_id
            WHERE a.contrato_id = %s
            ORDER BY a.criado_em
        """, (contrato_id,))
        assinaturas = [dict(r) for r in cur.fetchall()]

        return {**dict(contrato), "assinaturas": assinaturas}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()


# -------------------------------------------------------------
# GET /contratos/{id}/documento — gera o .docx do contrato sob demanda
# -------------------------------------------------------------

TIPO_LABELS_DOC = {
    "agricola": "Contrato de Parceria Agrícola",
    "pecuaria": "Contrato de Parceria Pecuária",
    "agroindustrial": "Contrato de Parceria Agroindustrial",
    "extrativa": "Contrato de Parceria Extrativa",
    "condominio": "Contrato de Condomínio Rural",
    "arrendamento": "Contrato de Arrendamento Rural",
    "comodato": "Contrato de Comodato Rural",
    "compra_venda": "Contrato de Compra e Venda Rural",
    "integracao_agroindustrial": "Contrato de Integração Agroindustrial",
}


def _fmt_data_doc(d) -> str:
    if not d:
        return "Prazo indeterminado"
    try:
        if isinstance(d, (datetime, date)):
            return d.strftime("%d/%m/%Y")
        return datetime.strptime(str(d)[:10], "%Y-%m-%d").strftime("%d/%m/%Y")
    except Exception:
        return str(d)


def _gerar_docx_contrato(c: dict) -> bytes:
    """Gera um .docx a partir dos dados de um contrato (linha de
    vw_contratos_resumo). Layout simples e genérico — serve como minuta
    inicial; não substitui revisão jurídica antes da assinatura formal."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    tipo = c.get("tipo") or ""
    titulo = TIPO_LABELS_DOC.get(tipo, f"Contrato Rural — {tipo}")

    h = doc.add_heading(titulo, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        f"Pelo presente instrumento particular, as partes abaixo qualificadas "
        f"celebram o presente {titulo.lower()}, que se regerá pelas cláusulas "
        f"e condições a seguir estabelecidas."
    )

    from docx.shared import RGBColor

    doc.add_heading("Partes", level=2)
    p_out = doc.add_paragraph()
    p_out.add_run("Outorgante: ").bold = True
    if c.get("outorgante_nome"):
        p_out.add_run(c["outorgante_nome"])
    else:
        r = p_out.add_run("[DADO PENDENTE — CADASTRAR OUTORGANTE ANTES DA ASSINATURA]")
        r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        r.bold = True
    p_outd = doc.add_paragraph()
    p_outd.add_run("Outorgado: ").bold = True
    if c.get("outorgado_nome"):
        p_outd.add_run(c["outorgado_nome"])
    else:
        r = p_outd.add_run("[DADO PENDENTE — CADASTRAR OUTORGADO ANTES DA ASSINATURA]")
        r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        r.bold = True

    doc.add_heading("Objeto e Condições", level=2)
    tabela = doc.add_table(rows=0, cols=2)
    tabela.style = "Light Grid Accent 1"
    tabela.autofit = True

    def add_row(label, valor):
        cells = tabela.add_row().cells
        cells[0].text = label
        cells[1].text = str(valor) if valor not in (None, "") else "—"

    add_row("Data de início", _fmt_data_doc(c.get("data_inicio")))
    add_row("Data de término", _fmt_data_doc(c.get("data_fim")))

    _clausulas_raw = c.get("clausulas_adicionais")
    if isinstance(_clausulas_raw, str):
        try:
            _clausulas_raw = json.loads(_clausulas_raw)
        except Exception:
            _clausulas_raw = {}
    _aviso_previo = (_clausulas_raw or {}).get("aviso_previo_rescisao_dias")
    if not c.get("data_fim") and _aviso_previo:
        add_row("Aviso prévio para rescisão", f"{_aviso_previo} dias")

    if c.get("percentual_outorgante") is not None:
        add_row("Percentual do outorgante", f"{c['percentual_outorgante']}%")
        add_row("Percentual do outorgado", f"{c['percentual_outorgado']}%")

    if c.get("area_parceria_hectares") is not None:
        add_row("Área objeto do contrato", f"{float(c['area_parceria_hectares']):.2f} ha")

    if c.get("frequencia_pagamento"):
        FREQ_LABELS = {
            "safra": "Por safra", "mensal": "Mensal", "anual": "Anual",
            "semestral": "Semestral", "apos_abate": "Após abate",
            "ao_termino": "Ao término do contrato",
        }
        freq_bruta = c["frequencia_pagamento"]
        add_row("Frequência de pagamento", FREQ_LABELS.get(freq_bruta, freq_bruta))

    status_bruto = c.get("status") or "rascunho"
    status_label_doc = {"rascunho": "Pendente de Assinatura"}.get(status_bruto, status_bruto.replace("_", " ").title())
    add_row("Status atual", status_label_doc)

    clausulas = c.get("clausulas_adicionais")
    if clausulas:
        if isinstance(clausulas, str):
            try:
                clausulas = json.loads(clausulas)
            except Exception:
                clausulas = {"Observações": clausulas}
        ROTULOS_CLAUSULAS = {
            "quantidade_animais": "Quantidade de animais",
            "valor_investido_outorgante": "Valor investido na aquisição — Outorgante",
            "valor_investido_outorgado": "Valor investido na aquisição — Outorgado",
            "modalidade_parceria": "Modalidade da parceria",
            "especie_raca": "Espécie / Raça",
            "peso_medio_entrada_kg": "Peso médio de entrada (kg)",
        }
        CHAVES_JA_MOSTRADAS = {"aviso_previo_rescisao_dias", "responsabilidade_custos", "responsabilidade_riscos"}
        MODALIDADE_LABELS = {
            "pastagem": "Parceria de pastagem", "confinamento": "Confinamento",
            "integracao": "Integração (área + instalação + animais)",
        }
        CHAVES_MOEDA = {"valor_investido_outorgante", "valor_investido_outorgado"}
        if isinstance(clausulas, dict) and clausulas:
            itens_visiveis = {
                k: v for k, v in clausulas.items()
                if k not in CHAVES_JA_MOSTRADAS and v not in (None, "")
            }
            if itens_visiveis:
                doc.add_heading("Cláusulas Adicionais", level=2)
                for chave, valor in itens_visiveis.items():
                    rotulo = ROTULOS_CLAUSULAS.get(chave, chave.replace("_", " ").capitalize())
                    if chave == "modalidade_parceria":
                        valor = MODALIDADE_LABELS.get(str(valor), str(valor))
                    elif chave in CHAVES_MOEDA:
                        try:
                            valor = f"R$ {float(valor):,.2f}".replace(",", "@").replace(".", ",").replace("@", ".")
                        except (TypeError, ValueError):
                            pass
                    p = doc.add_paragraph()
                    p.add_run(f"{rotulo}: ").bold = True
                    p.add_run(str(valor))

    if not c.get("data_fim") and _aviso_previo:
        doc.add_heading("Da Vigência e Rescisão", level=2)
        doc.add_paragraph(
            f"O prazo de vigência deste contrato é indeterminado, podendo qualquer "
            f"das partes denunciá-lo mediante comunicação por escrito à outra parte, "
            f"com antecedência mínima de {_aviso_previo} dias."
        )

    freq_bruta_final = c.get("frequencia_pagamento")
    if freq_bruta_final == "safra":
        doc.add_paragraph(
            "Entende-se como \"safra\", para fins deste contrato, o ciclo completo "
            "de manejo do rebanho até o momento definido para apuração e divisão "
            "dos resultados entre as partes."
        )

    custos_operacionais = (_clausulas_raw or {}).get("responsabilidade_custos")
    if custos_operacionais:
        doc.add_heading("Dos Custos e Despesas Operacionais", level=2)
        doc.add_paragraph(str(custos_operacionais))

    riscos_perdas = (_clausulas_raw or {}).get("responsabilidade_riscos")
    if riscos_perdas:
        doc.add_heading("Dos Riscos e Perdas", level=2)
        doc.add_paragraph(str(riscos_perdas))

    doc.add_heading("Assinaturas", level=2)
    doc.add_paragraph()
    doc.add_paragraph("_" * 40)
    doc.add_paragraph(f"Outorgante: {c.get('outorgante_nome') or ''}")
    doc.add_paragraph()
    doc.add_paragraph("_" * 40)
    doc.add_paragraph(f"Outorgado: {c.get('outorgado_nome') or ''}")

    aviso = doc.add_paragraph()
    aviso.add_run(
        "Documento gerado automaticamente pelo RuralCaixa a partir dos dados "
        "cadastrados. Recomenda-se revisão antes da assinatura formal."
    ).italic = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@router.get("/{contrato_id}/documento")
def gerar_documento_contrato(contrato_id: str):
    if contrato_id in _STATIC_PATHS:
        raise HTTPException(status_code=404, detail="Rota de frontend — não é um contrato.")
    conn = get_db()
    try:
        cur = conn.cursor()
        cur.execute("SELECT * FROM vw_contratos_resumo WHERE id = %s", (contrato_id,))
        contrato = cur.fetchone()
        if not contrato:
            raise HTTPException(status_code=404, detail="Contrato não encontrado")
        contrato = dict(contrato)

        # vw_contratos_resumo não expõe clausulas_adicionais — busca direto da tabela
        cur.execute("SELECT clausulas_adicionais FROM contratos WHERE id = %s", (contrato_id,))
        row_extra = cur.fetchone()
        if row_extra:
            contrato["clausulas_adicionais"] = row_extra["clausulas_adicionais"]

        docx_bytes = _gerar_docx_contrato(contrato)

        tipo_slug = (contrato.get("tipo") or "contrato").replace(" ", "_")
        nome_arquivo = f"contrato_{tipo_slug}_{contrato_id[:8]}.docx"

        return Response(
            content=docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{nome_arquivo}"'},
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Erro ao gerar documento: {e}")
    finally:
        conn.close()


# -------------------------------------------------------------
# POST /contratos
# -------------------------------------------------------------
@router.post("/", status_code=201)
def criar_contrato(body: ContratoCreate, request: Request):
    conn = get_db()
    try:
        cur = conn.cursor()

        # Resolver parceiro externo outorgante
        outorgante_externo_id = None
        if body.outorgante_externo:
            cur.execute("""
                INSERT INTO parceiros_externos (nome, tipo_documento, documento, telefone, email)
                VALUES (%s, %s, %s, %s, %s)
                ON CONFLICT (tipo_documento, documento)
                DO UPDATE SET nome = EXCLUDED.nome
                RETURNING id
            """, (
                body.outorgante_externo.nome,
                body.outorgante_externo.tipo_documento,
                body.outorgante_externo.documento,
                body.outorgante_externo.telefone,
                body.outorgante_externo.email,
            ))
            outorgante_externo_id = cur.fetchone()["id"]

        # Resolver parceiro externo outorgado
        outorgado_externo_id = None
        if body.outorgado_externo:
            cur.execute("""
                INSERT INTO parceiros_externos (nome, tipo_documento, documento, telefone, email)
                VALUES (%s, %s, %s, %s, %s)
                ON CONFLICT (tipo_documento, documento)
                DO UPDATE SET nome = EXCLUDED.nome
                RETURNING id
            """, (
                body.outorgado_externo.nome,
                body.outorgado_externo.tipo_documento,
                body.outorgado_externo.documento,
                body.outorgado_externo.telefone,
                body.outorgado_externo.email,
            ))
            outorgado_externo_id = cur.fetchone()["id"]

        # Inserir contrato
        cur.execute("""
            INSERT INTO contratos (
                fazenda_id, tipo, status,
                outorgante_socio_id, outorgante_externo_id,
                outorgado_socio_id,  outorgado_externo_id,
                data_inicio, data_fim,
                percentual_outorgante, percentual_outorgado,
                frequencia_pagamento, area_parceria_hectares,
                clausulas_adicionais
            ) VALUES (%s,%s,'rascunho',%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)
            RETURNING *
        """, (
            body.fazenda_id, body.tipo,
            body.outorgante_socio_id, outorgante_externo_id,
            body.outorgado_socio_id,  outorgado_externo_id,
            body.data_inicio, body.data_fim,
            body.percentual_outorgante, body.percentual_outorgado,
            body.frequencia_pagamento, body.area_parceria_hectares,
            json.dumps(body.clausulas_adicionais or {}),
        ))
        contrato = dict(cur.fetchone())

        log_auditoria(cur, contrato["id"], "contrato_criado",
                      f"Contrato {body.tipo} criado", str(request.client.host))

        conn.commit()
        return {"data": contrato}
    except HTTPException:
        raise
    except Exception as e:
        conn.rollback()
        err_str = str(e)
        if "foreign key" in err_str.lower() or "violates" in err_str.lower():
            raise HTTPException(
                status_code=422,
                detail=f"fazenda_id inválido ou não encontrado: {body.fazenda_id}. "
                       f"Verifique se a propriedade existe antes de criar o contrato."
            )
        raise HTTPException(status_code=500, detail=err_str)
    finally:
        conn.close()


# -------------------------------------------------------------
# POST /contratos/{id}/enviar
# -------------------------------------------------------------
@router.post("/{contrato_id}/enviar")
def enviar_para_assinatura(contrato_id: str, request: Request):
    if contrato_id in _STATIC_PATHS:
        raise HTTPException(status_code=404, detail="Rota de frontend — não é um contrato.")
    conn = get_db()
    try:
        cur = conn.cursor()

        cur.execute("SELECT * FROM contratos WHERE id = %s", (contrato_id,))
        contrato = cur.fetchone()
        if not contrato:
            raise HTTPException(status_code=404, detail="Contrato não encontrado")
        if contrato["status"] != "rascunho":
            raise HTTPException(status_code=400,
                detail=f"Contrato já está em status '{contrato['status']}'")

        cur.execute("""
            UPDATE contratos SET status = 'aguardando_assinaturas', atualizado_em = NOW()
            WHERE id = %s
        """, (contrato_id,))

        partes = _resolver_partes(cur, contrato)
        partes_notificadas = []

        import os
        frontend_url = os.getenv("FRONTEND_URL", "https://ruralcaixa-mvp.vercel.app")

        for parte in partes:
            otp = gerar_otp()
            otp_hash = hash_otp(otp)
            expira = datetime.now() + timedelta(minutes=30)
            link = f"{frontend_url}/assinar/{contrato_id}?parte={parte['papel']}"

            cur.execute("""
                INSERT INTO assinaturas (
                    contrato_id, papel, socio_id, parceiro_externo_id,
                    token_otp, token_expira_em, link_enviado_em, pdf_hash_no_momento
                ) VALUES (%s,%s,%s,%s,%s,%s,NOW(),%s)
                RETURNING id
            """, (
                contrato_id, parte["papel"],
                parte.get("socio_id"), parte.get("parceiro_externo_id"),
                otp_hash, expira,
                contrato.get("pdf_hash_sha256"),
            ))
            assinatura_id = cur.fetchone()["id"]

            def _mascarar_telefone(tel: str) -> str:
                if not tel:
                    return "número cadastrado"
                t = tel.replace(" ", "").replace("-", "")
                if len(t) >= 6:
                    return t[:4] + "•" * (len(t) - 6) + t[-2:]
                return "•" * len(t)

            if parte.get("telefone"):
                _enviar_whatsapp_otp(parte["telefone"], parte["nome"], otp, link)

            log_auditoria(cur, contrato_id, "link_assinatura_enviado",
                         f"Link enviado para {parte['nome']} ({parte['papel']})",
                         str(request.client.host))

            partes_notificadas.append({
                "papel": parte["papel"],
                "nome": parte["nome"],
                "assinatura_id": str(assinatura_id),
                "whatsapp_enviado": bool(parte.get("telefone")),
                "telefone_mascarado": _mascarar_telefone(parte.get("telefone", ""))
            })

        conn.commit()
        return {"message": "Contrato enviado para assinatura",
                "partes_notificadas": partes_notificadas}
    except HTTPException:
        raise
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()


# -------------------------------------------------------------
# POST /contratos/{id}/assinar
# -------------------------------------------------------------
@router.post("/{contrato_id}/assinar")
def assinar_contrato(contrato_id: str, body: AssinarRequest, request: Request):
    if contrato_id in _STATIC_PATHS:
        raise HTTPException(status_code=404, detail="Rota de frontend — não é um contrato.")
    conn = get_db()
    try:
        cur = conn.cursor()

        cur.execute("""
            SELECT * FROM assinaturas
            WHERE contrato_id = %s AND papel = %s AND status IN ('pendente','visualizado')
        """, (contrato_id, body.papel))
        assinatura = cur.fetchone()

        if not assinatura:
            raise HTTPException(status_code=404,
                detail="Assinatura não encontrada ou já concluída")

        if datetime.now() > assinatura["token_expira_em"].replace(tzinfo=None):
            log_auditoria(cur, contrato_id, "otp_expirado", ip=str(request.client.host))
            conn.commit()
            raise HTTPException(status_code=400, detail="OTP expirado. Solicite novo envio.")

        if assinatura["token_tentativas"] >= 5:
            raise HTTPException(status_code=429, detail="Muitas tentativas. Solicite novo OTP.")

        otp_valido = hash_otp(body.otp) == assinatura["token_otp"]

        if not otp_valido:
            cur.execute("UPDATE assinaturas SET token_tentativas = token_tentativas + 1 WHERE id = %s",
                        (assinatura["id"],))
            log_auditoria(cur, contrato_id, "otp_falhou", ip=str(request.client.host))
            conn.commit()
            raise HTTPException(status_code=400, detail="OTP inválido")

        cur.execute("""
            UPDATE assinaturas SET
                status = 'assinado',
                assinado_em = NOW(),
                ip_assinatura = %s,
                user_agent = %s,
                geolocalizacao = %s,
                token_otp = NULL,
                token_tentativas = token_tentativas + 1
            WHERE id = %s
        """, (
            str(request.client.host),
            request.headers.get("user-agent", ""),
            json.dumps(body.geolocalizacao or {}),
            assinatura["id"],
        ))

        log_auditoria(cur, contrato_id, "contrato_assinado",
                     f"Assinatura de {body.papel} registrada",
                     str(request.client.host),
                     {"geolocalizacao": body.geolocalizacao})

        cur.execute("""
            SELECT COUNT(*) AS total,
                   SUM(CASE WHEN status = 'assinado' THEN 1 ELSE 0 END) AS assinadas
            FROM assinaturas WHERE contrato_id = %s
        """, (contrato_id,))
        counts = cur.fetchone()

        if counts["total"] > 0 and counts["total"] == counts["assinadas"]:
            cur.execute("""
                UPDATE contratos SET status = 'ativo', atualizado_em = NOW()
                WHERE id = %s AND status = 'aguardando_assinaturas'
            """, (contrato_id,))
            log_auditoria(cur, contrato_id, "contrato_ativado",
                         "Todas as partes assinaram. Contrato ativado.")

        conn.commit()
        return {"message": "Assinatura registrada com sucesso"}
    except HTTPException:
        raise
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()


# -------------------------------------------------------------
# GET /contratos/{id}/auditoria
# -------------------------------------------------------------
@router.get("/{contrato_id}/auditoria")
def auditoria_contrato(contrato_id: str):
    if contrato_id in _STATIC_PATHS:
        raise HTTPException(status_code=404, detail="Rota de frontend — não é um contrato.")
    conn = get_db()
    try:
        cur = conn.cursor()
        cur.execute("""
            SELECT id, evento, descricao, ip, metadata, criado_em
            FROM auditoria_contratos
            WHERE contrato_id = %s
            ORDER BY criado_em ASC
        """, (contrato_id,))
        rows = [dict(r) for r in cur.fetchall()]
        return {"data": rows}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()


# -------------------------------------------------------------
# DELETE /contratos/{id}
# -------------------------------------------------------------
@router.delete("/{contrato_id}")
def deletar_contrato(contrato_id: str):
    if contrato_id in _STATIC_PATHS:
        raise HTTPException(status_code=404, detail="Rota de frontend — não é um contrato.")
    conn = get_db()
    try:
        cur = conn.cursor()
        cur.execute("SELECT status FROM contratos WHERE id = %s", (contrato_id,))
        row = cur.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Não encontrado")
        if row["status"] != "rascunho":
            raise HTTPException(status_code=400,
                detail="Apenas rascunhos podem ser excluídos")
        cur.execute("DELETE FROM contratos WHERE id = %s", (contrato_id,))
        conn.commit()
        return {"message": "Contrato excluído"}
    except HTTPException:
        raise
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()


# -------------------------------------------------------------
# POST /contratos/{id}/condominos
# ✅ CORREÇÃO 4: usa área em hectares — percentual calculado pelo sistema
# -------------------------------------------------------------
# NOTA: os endpoints de condomínio (adicionar/listar condôminos) foram
# removidos daqui — essa funcionalidade foi migrada para o módulo dedicado
# app/routers/condominio.py (prefixo /condominio), que usa sua própria
# tabela (condominio_condominos) e fluxo de assinatura por condômino via OTP.


# -------------------------------------------------------------
# HELPERS INTERNOS
# -------------------------------------------------------------

def _resolver_partes(cur, contrato):
    partes = []

    if contrato["outorgante_socio_id"]:
        cur.execute("SELECT id, nome, telefone FROM produtores WHERE id = %s",
                    (contrato["outorgante_socio_id"],))
        s = cur.fetchone()
        if s:
            partes.append({"papel": "outorgante", "socio_id": s["id"],
                           "nome": s["nome"], "telefone": s.get("telefone")})
    elif contrato["outorgante_externo_id"]:
        cur.execute("SELECT id, nome, telefone FROM parceiros_externos WHERE id = %s",
                    (contrato["outorgante_externo_id"],))
        pe = cur.fetchone()
        if pe:
            partes.append({"papel": "outorgante", "parceiro_externo_id": pe["id"],
                           "nome": pe["nome"], "telefone": pe.get("telefone")})

    if contrato["outorgado_socio_id"]:
        cur.execute("SELECT id, nome, telefone FROM produtores WHERE id = %s",
                    (contrato["outorgado_socio_id"],))
        s = cur.fetchone()
        if s:
            partes.append({"papel": "outorgado", "socio_id": s["id"],
                           "nome": s["nome"], "telefone": s.get("telefone")})
    elif contrato["outorgado_externo_id"]:
        cur.execute("SELECT id, nome, telefone FROM parceiros_externos WHERE id = %s",
                    (contrato["outorgado_externo_id"],))
        pe = cur.fetchone()
        if pe:
            partes.append({"papel": "outorgado", "parceiro_externo_id": pe["id"],
                           "nome": pe["nome"], "telefone": pe.get("telefone")})

    return partes


def _enviar_whatsapp_otp(telefone: str, nome: str, otp: str, link: str):
    import os, requests
    phone_id = os.getenv("WHATSAPP_PHONE_ID", "1154361321082939")
    token    = os.getenv("WHATSAPP_TOKEN", "")
    if not token:
        print(f"[WARN] WHATSAPP_TOKEN não configurado. OTP para {nome}: {otp}")
        return None
    url = f"https://graph.facebook.com/v19.0/{phone_id}/messages"
    payload = {
        "messaging_product": "whatsapp",
        "to": telefone,
        "type": "template",
        "template": {
            "name": "assinatura_contrato",
            "language": {"code": "pt_BR"},
            "components": [
                {"type": "body", "parameters": [{"type": "text", "text": otp}]},
                {"type": "button", "sub_type": "url", "index": "0",
                 "parameters": [{"type": "text", "text": otp}]}
            ]
        }
    }
    headers = {"Authorization": f"Bearer {token}", "Content-Type": "application/json"}
    try:
        r = requests.post(url, json=payload, headers=headers, timeout=10)
        print(f"[WhatsApp] Status: {r.status_code} Resposta: {r.text}")
        return r.json().get("messages", [{}])[0].get("id")
    except Exception as e:
        print(f"[WARN] Erro WhatsApp para {telefone}: {e}")
        return None
