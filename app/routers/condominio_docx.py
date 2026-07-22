"""
condominio_docx.py — Gerador de contrato DOCX para Condomínio Rural
Endpoint: GET /condominio/{contrato_id}/documento
"""
from fastapi import APIRouter, HTTPException, Request
from fastapi.responses import StreamingResponse
from datetime import date, datetime
from decimal import Decimal
import io, os, psycopg2, psycopg2.extras

from docx import Document as DocxDocument
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

router = APIRouter(prefix="/condominio", tags=["Condomínio Rural"])

DB_URL = os.getenv("DATABASE_URL", "postgresql://postgres:tkyfcRsbrZuuHoThKgjuTiZWYVXOTdOX@gondola.proxy.rlwy.net:53900/railway")

def get_db():
    return psycopg2.connect(DB_URL, cursor_factory=psycopg2.extras.RealDictCursor)

def fmtdata(d):
    if not d: return "___/___/______"
    if isinstance(d, (date, datetime)):
        return d.strftime("%d/%m/%Y")
    return str(d)

def fmtha(v):
    if v is None: return "0,0000"
    return f"{float(v):,.4f}".replace(",", "X").replace(".", ",").replace("X", ".")

def fmtpct(v):
    if v is None: return "0,00%"
    return f"{float(v):.2f}%".replace(".", ",")

def anos_entre(d1, d2):
    if not d1 or not d2: return 0
    anos = (d2 - d1).days / 365.25
    return round(anos, 1)

EXTENSO = {
    1: "um", 2: "dois", 3: "três", 4: "quatro", 5: "cinco",
    6: "seis", 7: "sete", 8: "oito", 9: "nove", 10: "dez",
    15: "quinze", 20: "vinte", 25: "vinte e cinco", 30: "trinta"
}

def num_extenso(n):
    return EXTENSO.get(int(n), str(n))

def buscar_dados_contrato(contrato_id: str) -> dict:
    conn = get_db()
    try:
        cur = conn.cursor()

        # Contrato principal
        cur.execute("""
            SELECT c.*, i.nome AS imovel_nome, i.nirf AS matricula,
                   i.municipio, i.uf, i.area_ha AS area_total_imovel_ha
            FROM contratos c
            JOIN imoveis_rurais i ON i.id = c.fazenda_id
            WHERE c.id = %s AND c.tipo = 'condominio'
        """, (contrato_id,))
        contrato = cur.fetchone()
        if not contrato:
            raise HTTPException(404, "Contrato de condomínio não encontrado")

        # Condôminos
        cur.execute("""
            SELECT cc.*,
                   p.nome  AS prod_nome,  p.cpf  AS prod_cpf,
                   pe.nome AS parc_nome, pe.documento AS parc_doc,
                   pe.tipo_documento AS parc_tipo_doc
            FROM contrato_condominos cc
            LEFT JOIN produtores p ON p.id = cc.produtor_id
            LEFT JOIN parceiros_externos pe ON pe.id = cc.parceiro_id
            WHERE cc.contrato_id = %s AND cc.ativo = TRUE
            ORDER BY cc.criado_em
        """, (contrato_id,))
        condominos = [dict(r) for r in cur.fetchall()]

        return {"contrato": dict(contrato), "condominos": condominos}
    finally:
        conn.close()

def montar_variaveis(dados: dict) -> dict:
    c = dados["contrato"]
    conds = dados["condominos"]
    area_ind = float(c.get("area_parceria_hectares") or 25)
    area_tot = float(c.get("area_total_imovel_ha") or 130)
    area_com = round(area_tot - area_ind, 4)
    anos = anos_entre(c.get("data_inicio"), c.get("data_fim"))

    vars = {
        "nome_imovel": c.get("imovel_nome") or "{{nome_imovel}}",
        "matricula_imovel": c.get("matricula") or "{{matricula_imovel}}",
        "municipio_imovel": c.get("municipio") or "{{municipio_imovel}}",
        "uf_imovel": c.get("uf") or "{{uf_imovel}}",
        "area_total_imovel_ha": fmtha(area_tot),
        "area_total_ha": fmtha(area_ind),
        "area_comum_ha": fmtha(area_com),
        "data_inicio": fmtdata(c.get("data_inicio")),
        "data_fim": fmtdata(c.get("data_fim")),
        "prazo_vigencia_anos": str(int(anos)) if anos == int(anos) else str(anos),
        "prazo_vigencia_anos_extenso": num_extenso(int(anos)) if anos >= 1 else str(anos),
        "tipo_exploracao": "pecuária e agricultura",
        "data_assinatura": fmtdata(date.today()),
        # Governança — valores padrão (editáveis no Word)
        "nome_administrador": "{{nome_administrador}}",
        "percentual_administrador": "{{percentual_administrador}}",
        "prazo_mandato_administrador": "2",
        "limite_contrato_administrador": "R$ 5.000,00",
        "dia_distribuicao": "15",
        "prazo_distribuicao_dias": "30",
        "remuneracao_administrador": "{{remuneracao_administrador}}",
        "rubrica_remuneracao": "pró-labore",
        "mes_assembleia_ordinaria": "março",
        "tolerancia_desvio_plano_pct": "20",
        # Parcerias
        "limite_parceria_interna_pct": "50",
        "limite_concentracao_pct": "70",
        "limite_parceria_terceiros_pct": "30",
        "prazo_max_parceria_terceiros": "5",
        # Retirada
        "prazo_notificacao_retirada_meses": "6 (seis)",
        "prazo_pagamento_retirada_meses": "12 (doze)",
        # Ambiental
        "percentual_rl": "20",
        # Conflitos
        "orgao_mediacao": "Câmara de Mediação e Arbitragem Rural",
        "foro_eleito": c.get("municipio") or "{{municipio_imovel}}",
        # Penalidades
        "valor_multa_leve": "R$ 500,00",
        "valor_multa_grave": "R$ 2.000,00",
        "limite_benfeitoria_aprovacao": "R$ 10.000,00",
        # Aportes
        "aporte_condomino_1": "{{aporte_condomino_1}}",
        "aporte_condomino_2": "{{aporte_condomino_2}}",
        "aporte_condomino_3": "{{aporte_condomino_3}}",
    }

    # Condôminos
    for i, cond in enumerate(conds[:6], 1):
        nome = cond.get("prod_nome") or cond.get("parc_nome") or f"Condômino {i}"
        doc = cond.get("prod_cpf") or cond.get("parc_doc") or "{{doc}}"
        tipo_doc = "CPF" if cond.get("prod_cpf") else (cond.get("parc_tipo_doc") or "CPF/CNPJ")
        pct = float(cond.get("percentual_cota") or 0)
        area_cond = round(area_ind * pct / 100, 4)
        papel = cond.get("papel") or "condômino"

        vars[f"condomino_{i}_nome"] = nome
        vars[f"condomino_{i}_doc"] = f"{tipo_doc}: {doc}"
        vars[f"condomino_{i}_percentual"] = fmtpct(pct)
        vars[f"condomino_{i}_area_ha"] = fmtha(area_cond)
        vars[f"condomino_{i}_papel"] = papel.capitalize()

    # Preenche condôminos faltantes
    for i in range(len(conds) + 1, 7):
        vars[f"condomino_{i}_nome"] = ""
        vars[f"condomino_{i}_doc"] = ""
        vars[f"condomino_{i}_percentual"] = ""
        vars[f"condomino_{i}_area_ha"] = ""
        vars[f"condomino_{i}_papel"] = ""

    # Administrador = primeiro condômino com papel administrador
    adm = next((c for c in conds if "admin" in (c.get("papel") or "").lower()), conds[0] if conds else None)
    if adm:
        vars["nome_administrador"] = adm.get("prod_nome") or adm.get("parc_nome") or "{{nome_administrador}}"
        pct_adm = float(adm.get("percentual_cota") or 0)
        vars["percentual_administrador"] = fmtpct(pct_adm).replace("%", "").strip()

    return vars

def substituir_no_docx(template_bytes: bytes, variaveis: dict) -> bytes:
    """Substitui {{variavel}} em todos os parágrafos e células do documento."""
    doc = DocxDocument(io.BytesIO(template_bytes))

    def subst_paragrafo(para):
        """Substitui variáveis em um parágrafo, consolidando runs."""
        texto_completo = "".join(r.text for r in para.runs)
        novo_texto = texto_completo
        for k, v in variaveis.items():
            novo_texto = novo_texto.replace("{{" + k + "}}", str(v))

        if novo_texto != texto_completo and para.runs:
            # Preserva formatação do primeiro run, limpa os demais
            para.runs[0].text = novo_texto
            for r in para.runs[1:]:
                r.text = ""

    def processar_corpo(elemento):
        for para in elemento.paragraphs:
            subst_paragrafo(para)
        for tabela in elemento.tables:
            for row in tabela.rows:
                for cell in row.cells:
                    processar_corpo(cell)

    processar_corpo(doc)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


@router.get("/{contrato_id}/documento")
def gerar_documento(contrato_id: str, request: Request):
    """
    Gera o contrato DOCX de condomínio rural preenchido com os dados do banco.
    Retorna o arquivo para download.
    """
    dados = buscar_dados_contrato(contrato_id)
    variaveis = montar_variaveis(dados)

    # Carrega o template DOCX do disco
    template_path = os.path.join(
        os.path.dirname(__file__), "..", "templates", "contrato_condominio_rural.docx"
    )
    if not os.path.exists(template_path):
        raise HTTPException(
            500,
            f"Template não encontrado em {template_path}. "
            "Copie o arquivo contrato_condominio_rural.docx para app/templates/"
        )

    with open(template_path, "rb") as f:
        template_bytes = f.read()

    docx_bytes = substituir_no_docx(template_bytes, variaveis)

    nome_imovel = variaveis.get("nome_imovel", "condominio").replace(" ", "_")
    filename = f"Contrato_Condominio_{nome_imovel}_{date.today().strftime('%Y%m%d')}.docx"

    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )


@router.get("/{contrato_id}/documento/preview")
def preview_variaveis(contrato_id: str):
    """Retorna as variáveis que seriam substituídas no contrato (para debug)."""
    dados = buscar_dados_contrato(contrato_id)
    return montar_variaveis(dados)
