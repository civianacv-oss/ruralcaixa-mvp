"""
app/routers/recibos.py — RuralCaixa MVP

Módulo de Recibos com assinatura via WhatsApp (OTP), similar ao fluxo de
assinatura de Contratos Rurais / Condomínio. Ao ser assinado, o recibo é
acoplado a um lançamento financeiro (existente, se informado na criação, ou
um novo lançamento DESPESA criado automaticamente).
"""

import os
import time
import random
import hashlib
import logging
import io
from datetime import datetime, timedelta
from typing import Optional

import httpx
from fastapi import APIRouter, HTTPException, Request
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field

from docx import Document as DocxDocument
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

router = APIRouter(prefix="/recibos", tags=["Recibos"])
logger = logging.getLogger(__name__)

# WhatsApp Cloud API - mesmas variaveis usadas no restante do app
WAPP_TOKEN = os.getenv("WHATSAPP_TOKEN")
PHONE_ID = os.getenv("WHATSAPP_PHONE_ID")
GRAPH = "https://graph.facebook.com/v23.0"

# Rate limiting de OTP (max 3 por recibo por hora) — mesmo padrao do Condominio
_otp_attempts: dict = {}
_MAX_OTP_PER_HOUR = 3


def _check_otp_rate_limit(recibo_id: str) -> bool:
    now = time.time()
    tentativas = _otp_attempts.setdefault(recibo_id, [])
    tentativas[:] = [t for t in tentativas if now - t < 3600]
    if len(tentativas) >= _MAX_OTP_PER_HOUR:
        return False
    tentativas.append(now)
    return True


# ── Helpers ───────────────────────────────────────────────────────────────────

def get_db():
    import psycopg2
    import psycopg2.extras
    return psycopg2.connect(
        os.getenv("DATABASE_URL"),
        cursor_factory=psycopg2.extras.RealDictCursor,
    )


def gerar_otp() -> str:
    return f"{random.randint(0, 999999):06d}"


def hash_otp(otp: str) -> str:
    return hashlib.sha256(otp.encode()).hexdigest()


def _enviar_contexto_whatsapp(telefone: str, destinatario_nome: str, emissor_nome: str, valor: float, objeto: str) -> tuple[bool, str]:
    """
    Envia uma mensagem de contexto (categoria UTILITY) com os dados do recibo,
    antes do codigo de verificacao. Sem isso, a pessoa recebe so um codigo sem
    saber o que esta confirmando (templates AUTHENTICATION nao podem ter
    conteudo alem do codigo).
    """
    if not WAPP_TOKEN or not PHONE_ID:
        return False, "WHATSAPP_TOKEN/WHATSAPP_PHONE_ID nao configurado"
    if not telefone:
        return False, "Destinatario sem telefone cadastrado"

    valor_fmt = f"{valor:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    payload = {
        "messaging_product": "whatsapp",
        "recipient_type": "individual",
        "to": telefone,
        "type": "template",
        "template": {
            "name": "recibo_confirmacao_v2",
            "language": {"code": "pt_BR"},
            "components": [
                {
                    "type": "body",
                    "parameters": [
                        {"type": "text", "text": destinatario_nome},
                        {"type": "text", "text": emissor_nome},
                        {"type": "text", "text": valor_fmt},
                        {"type": "text", "text": objeto[:60]},
                    ],
                },
            ],
        },
    }
    try:
        with httpx.Client(timeout=10) as client:
            resp = client.post(
                f"{GRAPH}/{PHONE_ID}/messages",
                headers={"Authorization": f"Bearer {WAPP_TOKEN}", "Content-Type": "application/json"},
                json=payload,
            )
        data = resp.json()
        if "error" in data:
            erro = data["error"].get("message", str(data["error"]))
            logger.error(f"Erro ao enviar template de contexto do recibo: {erro}")
            return False, erro
        return True, "enviado"
    except Exception as e:
        logger.error(f"Excecao ao enviar contexto do recibo via WhatsApp: {e}")
        return False, str(e)


def _enviar_otp_whatsapp(telefone: str, otp: str) -> tuple[bool, str]:
    """Envia o OTP via template assinatura_contrato (categoria AUTHENTICATION)."""
    if not WAPP_TOKEN or not PHONE_ID:
        msg = "WHATSAPP_TOKEN/WHATSAPP_PHONE_ID nao configurado no ambiente"
        logger.warning(msg)
        return False, msg
    if not telefone:
        msg = "Destinatario sem telefone cadastrado"
        logger.warning(msg)
        return False, msg

    payload = {
        "messaging_product": "whatsapp",
        "recipient_type": "individual",
        "to": telefone,
        "type": "template",
        "template": {
            "name": "assinatura_contrato",
            "language": {"code": "pt_BR"},
            "components": [
                {"type": "body", "parameters": [{"type": "text", "text": otp}]},
                {"type": "button", "sub_type": "url", "index": "0",
                 "parameters": [{"type": "text", "text": otp}]},
            ],
        },
    }

    try:
        with httpx.Client(timeout=10) as client:
            resp = client.post(
                f"{GRAPH}/{PHONE_ID}/messages",
                headers={"Authorization": f"Bearer {WAPP_TOKEN}", "Content-Type": "application/json"},
                json=payload,
            )
        data = resp.json()
        if "error" in data:
            erro = data["error"].get("message", str(data["error"]))
            logger.error(f"Erro ao enviar template WhatsApp: {erro}")
            return False, erro
        return True, "enviado"
    except Exception as e:
        logger.error(f"Excecao ao enviar OTP via WhatsApp: {e}")
        return False, str(e)


def _gerar_docx_recibo(recibo: dict, produtor_nome: str) -> bytes:
    doc = DocxDocument()

    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run("RECIBO")
    run.bold = True
    run.font.size = Pt(20)

    doc.add_paragraph()

    valor_fmt = f"R$ {float(recibo['valor']):,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

    corpo = doc.add_paragraph()
    corpo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    corpo.add_run(
        f"Eu, {recibo['destinatario_nome']}, portador(a) do CPF/CNPJ nº "
        f"{recibo['destinatario_documento']}, declaro para os devidos fins ter "
        f"recebido de {produtor_nome} a quantia de {valor_fmt} "
        f"referente a: {recibo['objeto']}."
    )

    doc.add_paragraph()

    criado_em = recibo.get("criado_em")
    data_str = criado_em.strftime("%d/%m/%Y") if criado_em else datetime.now().strftime("%d/%m/%Y")
    doc.add_paragraph(f"Data de emissão: {data_str}")

    doc.add_paragraph()
    assinado_em = recibo.get("assinado_em")
    if assinado_em:
        doc.add_paragraph(
            f"Assinado digitalmente via WhatsApp em {assinado_em.strftime('%d/%m/%Y %H:%M')}, "
            f"com confirmação por código enviado ao telefone {recibo['destinatario_telefone']}."
        )
    else:
        doc.add_paragraph("(Documento ainda não assinado)")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ── Schemas ───────────────────────────────────────────────────────────────────

class ReciboCriar(BaseModel):
    destinatario_nome: str = Field(..., min_length=1)
    destinatario_documento: str = Field(..., min_length=1)
    destinatario_telefone: str = Field(..., min_length=8)
    objeto: str = Field(..., min_length=1)
    valor: float = Field(..., gt=0)
    lancamento_id: Optional[str] = None


class AssinarRecibo(BaseModel):
    otp: str


# ── Endpoints ─────────────────────────────────────────────────────────────────

@router.post("/", status_code=201)
def criar_recibo(body: ReciboCriar, request: Request):
    produtor_id = request.state.produtor_id
    conn = get_db()
    try:
        cur = conn.cursor()

        if body.lancamento_id:
            cur.execute("SELECT id FROM lancamentos WHERE id = %s AND produtor_id = %s",
                        (body.lancamento_id, produtor_id))
            if not cur.fetchone():
                raise HTTPException(status_code=404, detail="Lançamento não encontrado.")

        cur.execute("""
            INSERT INTO recibos (
                produtor_id, destinatario_nome, destinatario_documento,
                destinatario_telefone, objeto, valor, lancamento_id, status
            ) VALUES (%s, %s, %s, %s, %s, %s, %s, 'rascunho')
            RETURNING *
        """, (
            produtor_id, body.destinatario_nome, body.destinatario_documento,
            body.destinatario_telefone, body.objeto, body.valor, body.lancamento_id,
        ))
        recibo = cur.fetchone()
        conn.commit()
        return {"data": dict(recibo)}
    except HTTPException:
        raise
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()


@router.get("/")
def listar_recibos(request: Request):
    produtor_id = request.state.produtor_id
    conn = get_db()
    try:
        cur = conn.cursor()
        cur.execute("""
            SELECT * FROM recibos WHERE produtor_id = %s ORDER BY criado_em DESC
        """, (produtor_id,))
        return {"data": [dict(r) for r in cur.fetchall()]}
    finally:
        conn.close()


@router.get("/{recibo_id}")
def detalhe_recibo(recibo_id: str, request: Request):
    produtor_id = request.state.produtor_id
    conn = get_db()
    try:
        cur = conn.cursor()
        cur.execute("SELECT * FROM recibos WHERE id = %s AND produtor_id = %s",
                    (recibo_id, produtor_id))
        recibo = cur.fetchone()
        if not recibo:
            raise HTTPException(status_code=404, detail="Recibo não encontrado.")
        return {"data": dict(recibo)}
    finally:
        conn.close()


@router.delete("/{recibo_id}", status_code=204)
def excluir_recibo(recibo_id: str, request: Request):
    produtor_id = request.state.produtor_id
    conn = get_db()
    try:
        cur = conn.cursor()
        cur.execute(
            "DELETE FROM recibos WHERE id = %s AND produtor_id = %s AND status = 'rascunho' RETURNING id",
            (recibo_id, produtor_id)
        )
        if not cur.fetchone():
            raise HTTPException(
                status_code=404,
                detail="Recibo não encontrado ou não pode ser excluído (já enviado/assinado)."
            )
        conn.commit()
    except HTTPException:
        raise
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()


@router.post("/{recibo_id}/enviar-assinatura")
def enviar_assinatura_recibo(recibo_id: str, request: Request):
    produtor_id = request.state.produtor_id

    if not _check_otp_rate_limit(recibo_id):
        raise HTTPException(
            status_code=429,
            detail=f"Limite de {_MAX_OTP_PER_HOUR} OTPs por hora atingido para este recibo. Tente novamente mais tarde."
        )

    conn = get_db()
    try:
        cur = conn.cursor()
        cur.execute("""
            SELECT r.*, p.nome AS emissor_nome
            FROM recibos r
            JOIN produtores p ON p.id = r.produtor_id
            WHERE r.id = %s AND r.produtor_id = %s
        """, (recibo_id, produtor_id))
        recibo = cur.fetchone()
        if not recibo:
            raise HTTPException(status_code=404, detail="Recibo não encontrado.")
        if recibo["status"] == "assinado":
            raise HTTPException(status_code=400, detail="Recibo já foi assinado.")

        otp = gerar_otp()
        otp_hash_val = hash_otp(otp)
        expira = datetime.now() + timedelta(minutes=30)

        cur.execute("""
            UPDATE recibos SET status = 'aguardando_assinatura', otp_hash = %s,
                   otp_expira_em = %s, atualizado_em = NOW()
            WHERE id = %s
        """, (otp_hash_val, expira, recibo_id))
        conn.commit()

        _enviar_contexto_whatsapp(
            recibo["destinatario_telefone"], recibo["destinatario_nome"],
            recibo["emissor_nome"], float(recibo["valor"]), recibo["objeto"]
        )
        enviado, detalhe_envio = _enviar_otp_whatsapp(recibo["destinatario_telefone"], otp)

        response = {
            "message": "OTP gerado e enviado via WhatsApp." if enviado
                        else f"OTP gerado, mas o envio via WhatsApp falhou: {detalhe_envio}",
            "enviado_whatsapp": enviado,
            "expira_em": expira.isoformat(),
        }
        if os.getenv("DEBUG") == "true":
            response["otp_debug"] = otp
        return response
    except HTTPException:
        raise
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()


@router.post("/{recibo_id}/assinar")
def assinar_recibo(recibo_id: str, body: AssinarRecibo, request: Request):
    conn = get_db()
    try:
        cur = conn.cursor()
        cur.execute("SELECT * FROM recibos WHERE id = %s", (recibo_id,))
        recibo = cur.fetchone()
        if not recibo:
            raise HTTPException(status_code=404, detail="Recibo não encontrado.")
        if recibo["status"] == "assinado":
            raise HTTPException(status_code=400, detail="Recibo já foi assinado.")
        if not recibo["otp_hash"] or not recibo["otp_expira_em"]:
            raise HTTPException(status_code=400, detail="Nenhum código foi enviado para este recibo ainda.")
        if datetime.now() > recibo["otp_expira_em"]:
            raise HTTPException(status_code=400, detail="Código expirado. Solicite um novo.")
        if hash_otp(body.otp) != recibo["otp_hash"]:
            raise HTTPException(status_code=400, detail="Código inválido.")

        lancamento_id = recibo["lancamento_id"]

        if not lancamento_id:
            cur.execute(
                "INSERT INTO subcontas (nome, tipo, atividade_tipo) VALUES (%s, 'DESPESA', 'RURAL') RETURNING id",
                (recibo["objeto"][:255],)
            )
            subconta_id = cur.fetchone()["id"]

            cur.execute("""
                INSERT INTO lancamentos (produtor_id, subconta_id, valor, data, origem_modulo, origem_descricao)
                VALUES (%s, %s, %s, CURRENT_DATE, 'recibos', %s)
                RETURNING id
            """, (recibo["produtor_id"], subconta_id, recibo["valor"],
                  f"Recibo assinado: {recibo['objeto']}"))
            lancamento_id = cur.fetchone()["id"]
        else:
            cur.execute("""
                UPDATE lancamentos SET origem_modulo = 'recibos',
                       origem_descricao = %s
                WHERE id = %s
            """, (f"Confirmado por recibo assinado: {recibo['objeto']}", lancamento_id))

        cur.execute("""
            UPDATE recibos SET status = 'assinado', assinado_em = NOW(),
                   lancamento_id = %s, atualizado_em = NOW()
            WHERE id = %s
            RETURNING *
        """, (lancamento_id, recibo_id))
        atualizado = cur.fetchone()
        conn.commit()
        return {"data": dict(atualizado)}
    except HTTPException:
        raise
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()


@router.get("/{recibo_id}/documento")
def baixar_documento_recibo(recibo_id: str, request: Request):
    produtor_id = request.state.produtor_id
    conn = get_db()
    try:
        cur = conn.cursor()
        cur.execute("""
            SELECT r.*, p.nome AS produtor_nome
            FROM recibos r
            JOIN produtores p ON p.id = r.produtor_id
            WHERE r.id = %s AND r.produtor_id = %s
        """, (recibo_id, produtor_id))
        recibo = cur.fetchone()
        if not recibo:
            raise HTTPException(status_code=404, detail="Recibo não encontrado.")

        docx_bytes = _gerar_docx_recibo(dict(recibo), recibo["produtor_nome"])
        filename = f"Recibo_{recibo['destinatario_nome'].replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx"

        return StreamingResponse(
            io.BytesIO(docx_bytes),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    except HTTPException:
        raise
    finally:
        conn.close()
